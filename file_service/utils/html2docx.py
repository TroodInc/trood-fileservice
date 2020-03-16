import bs4
import tempfile
import cssutils
import re

from bs4 import BeautifulSoup
from collections import namedtuple
from docx import Document
from docx.shared import Cm, Pt

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TableCell = namedtuple('TableCell',
                       ['index', 'rowspan', 'colspan'])


def convert_px_to_pt(px):
    PX_TO_PT_COEFFICIENT = 0.75
    pt = px * PX_TO_PT_COEFFICIENT
    return pt


class TableNavigator:
    """
    Support class used to navigate on docx table.
    """
    def __init__(self, table):
        self.table = table
        self.navigation_array = self.get_array()

    def to_matrix(self, rows, cells_in_row):
        cells_count = rows * cells_in_row
        cells_idxs = list(range(cells_count))
        return [cells_idxs[i:i+cells_in_row] for i in range(0, cells_count, cells_in_row)] 

    def get_array(self):
        row_number = len(self.table.rows)
        column_number = len(self.table.columns)
        return self.to_matrix(row_number, column_number)

    def _get_row(self, index):
        for row_index, row in enumerate(self.navigation_array):
            if index in row:
                return row_index, row

    def _get_column(self, index):
        transpone_array = list(zip(*self.navigation_array))
        for column_index, column in enumerate(transpone_array):
            if index in column:
                return column_index, column


class TdThCell(TableNavigator):
    def __init__(self, table, tag, cells_in_tag):
        super().__init__(table)
        self.tag = tag
        self.cells_in_tag = cells_in_tag
        self.merge()

    @property
    def labled_cells(self):
        """
        Creates Labled_cells. Each labled_cell corresponds to docx cell by index.
        Labled_cells group coresponds to HTML tag. first  labled _cell contains
        information about rowspan and colspan.
        """
        labled_cells = []
        for num, cell_index in enumerate(sorted(self.cells_in_tag)):
            rowspan = None
            colspan = None
            if len(self.cells_in_tag) > 1 and num == 0:
                rowspan = int(self.tag.get('rowspan')) if self.tag.get('rowspan') else None
                colspan = int(self.tag.get('colspan')) if self.tag.get('colspan') else None

            labled_cell = TableCell(
                index=cell_index,
                rowspan=rowspan,
                colspan=colspan
            )
            labled_cells.append(labled_cell)
        return labled_cells

    @property
    def docx_cells(self):
        cells_indexes = {cell.index for cell in self.labled_cells}
        return [cell for idx, cell in enumerate(self.table._cells) if idx in cells_indexes]

    def merge(self):
        """
        Merge docx cells.
        """
        for cell in self.labled_cells:
            if cell.rowspan and not cell.colspan:
                self._merge_line(cell.index, cell.rowspan, vertical=True)
            elif cell.colspan and not cell.rowspan:
                self._merge_line(cell.index, cell.colspan, horizontal=True)
            elif cell.colspan and cell.rowspan:
                self._merge_rectangular(
                    cell.index, colspan=cell.colspan, rowspan=cell.rowspan
                    )

    def _merge_cells(self, first_cell_index, second_cell_index):
        first_cell = self.table._cells[first_cell_index]
        second_cell = self.table._cells[second_cell_index]
        first_cell.merge(second_cell)

    def _merge_line(self, cell_index, span, vertical=False, horizontal=False):
        if vertical and not horizontal:
            handler = self._get_column

        elif horizontal and not vertical:
            handler = self._get_row

        _, line = handler(cell_index)

        merging_cell_index = self._get_cell_index_in_array(cell_index, line,  span)
        self._merge_cells(cell_index, merging_cell_index)

    def _merge_rectangular(self, cell_index, colspan=None, rowspan=None):
        column_index, column = self._get_column(cell_index)
        initial_cell_index_in_column = column.index(cell_index)
        merging_cell_index = initial_cell_index_in_column + rowspan

        for cell_ind in column[initial_cell_index_in_column: merging_cell_index]:
            self._merge_line(cell_ind, colspan, horizontal=True)
        self._merge_line(cell_index, rowspan, vertical=True)

    def _get_cell_index_in_array(self, cell_index, line, span):
        cell_index_in_line = line.index(cell_index)
        span_shift = span - 1  # do not count cell itself
        index_in_array = line[cell_index_in_line + span_shift]
        return index_in_array


class TableCreator:
    """
    Creates empty table in docx document.
    """
    def __init__(self, html_table, document):
        self.html_table = html_table
        self.document = document
        self.row_number, self.column_number = self._get_table_shape()

    def create(self):
        return self.document.add_table(
            rows=self.row_number, cols=self.column_number, style='Table Grid'
            )

    def _get_table_shape(self):
        row_number = len(self.html_table.find_all('tr'))
        first_row = self.html_table.find_all('tr')[0]
        column_number = self._get_columns(first_row)
        return row_number, column_number

    def _get_columns(self, row):
        column_number = 0
        for table_cell in row.find_all(['td', 'th']):
            colspan = int(table_cell.get('colspan')) if table_cell.get('colspan') else None
            if colspan:
                column_number += colspan
            else:
                column_number += 1
        return column_number


class TdThCellCreator(TableNavigator):
    def __init__(self, table, tag, cell_index):
        """
        Creates marked_cells that corresponds to HTML tags.
        """
        super().__init__(table)
        self.tag = tag
        self.cell_index = cell_index
        self.cells_in_tag = self.gather_cells()

    def _get_cell_index_in_line(self, line, cell_index=None):
        if cell_index is None:
            cell_index = self.cell_index

        cell_index_in_line = line.index(cell_index)

        return cell_index_in_line

    def gather_cells(self):
        gathered_cells = []
        rowspan = int(self.tag.get('rowspan')) if self.tag.get('rowspan') else None
        colspan = int(self.tag.get('colspan')) if self.tag.get('colspan') else None

        if colspan and not rowspan:
            _, row = self._get_row(self.cell_index)
            index_in_line = self._get_cell_index_in_line(row)
            gathered_cells = row[index_in_line:index_in_line + colspan]
        elif rowspan and not colspan:
            _, column = self._get_column(self.cell_index)
            index_in_line = self._get_cell_index_in_line(column)
            gathered_cells = column[index_in_line:index_in_line + rowspan]

        elif rowspan and colspan:
            _, row = self._get_row(self.cell_index)
            index_in_line = self._get_cell_index_in_line(row)

            first_row = row[index_in_line:index_in_line + rowspan]

            for cell_ in first_row:
                _, column = self._get_column(cell_)
                index_in_line = self._get_cell_index_in_line(column, cell_index=cell_)

                gathered_cells.extend(column[index_in_line:index_in_line + colspan])

        else:
            gathered_cells.append(self.cell_index)

        return list(gathered_cells)

    def create_cell(self):
        return TdThCell(self.table, self.tag, self.cells_in_tag)


class TableDrawer(TableNavigator):
    """
    Convertes HTML table to docx table.
    Works as follows.
    1. Parse HTML table, count rows and columns in it.
    2. Creates empty docx table of size rows * columns.
    3. Lay out table into parts(marked_cells or TdThCell) that corresponds to HTML tags.
    For examle if rowspan=2 is presense in HTML tag marked part will consists of 2 docx cells.
    4. Merge docx cells in marked_cells
    5. Fill docx table with HTML tags content. Apply styles on it.
    """
    def __init__(self, html_table, document, raw_styles):
        self.raw_styles = raw_styles
        self.document = document
        self.html_table = html_table

        self.create_table()
        super().__init__(self.table)
        self.td_th_cells = self.mark_up_cells()

    def drop_cells(self, cells_to_drop):
        uprated_array = []
        for row in self.navigation_array:
            new_row = [cell for cell in row if cell not in cells_to_drop]
            uprated_array.append(new_row)
        return uprated_array

    def create_table(self):
        """
        Creates empty docx table of size rows * columns.
        """
        table_creator = TableCreator(self.html_table, self.document)
        self.table = table_creator.create()

    def mark_up_cells(self):
        """
        Lay out table into parts(marked_cells or TdThCell) that corresponds to HTML tags.
        """
        marked_cells = []
        for tr_index, tr in enumerate(self.html_table.find_all('tr')):
            index_row = self.navigation_array[tr_index]
            for tag_index, tag in enumerate(tr.find_all(['th', 'td'])):
                if tag_index == 0:
                    cell_index = index_row[0]

                cell_creator = TdThCellCreator(self.table, tag, cell_index)

                colspan = tag.get('colspan')
                if colspan:
                    cell_index += int(colspan)
                else:
                    cell_index += 1

                self.navigation_array = self.drop_cells(cell_creator.cells_in_tag)

                marked_cell = cell_creator.create_cell()
                marked_cells.append(marked_cell)

        return marked_cells

    def draw(self):
        """
        Fill docx table with HTML tags content. Apply styles on it.
        """
        for cell in self.td_th_cells:
            cell_styler = CellStyleProcesser(cell.tag, self.raw_styles, cell)
            cell_styler.apply()
            first_cell = cell.docx_cells[0]
            paragraph = first_cell.paragraphs[0]
            paragraph_styler = ParagraphStyleProcesser(cell.tag, self.raw_styles, paragraph)
            paragraph_styler.process()


class BaseStyleProcesser:
    def __init__(self, css_styles, tag=None):
        self.css_styles = css_styles
        if tag is not None:
            self.tag = tag
            self.get_all_rules()

        self.page_styles = self.get_page_rules()

    @property
    def rules(self):
        return cssutils.parseString(self.css_styles)

    def get_page_rules(self):
        def is_page_or_star(r):
            return r.type == r.PAGE_RULE or r.selectorText == "*"

        rules = [r for r in self.rules if is_page_or_star(r)]
        return rules

    def get_all_rules(self):
        """
        Parses selectors of following view:
        single selectors (td, th, table, tr, .className)
        complex selectors (.title td, .productInfo tfoot td, .productInfo .th1).
        TODO
        Does not parse selectors listed with a comma(.productInfo thead, .productInfo tbody)
        Does not parse id selectors( #timeEntry)
        Does not parse child selectors  (th:first-child)
        """
        self.tag_rules = set()
        tag_selectors = self.get_selectors(self.tag)
        for selector in tag_selectors:
            for rule in self.rules:
                if f"{selector}" == rule.selectorText:
                    self.tag_rules.add(rule)
                if f"{selector}" in rule.selectorText:
                    parent = self.tag.parent
                    wile_loop_limit = 0
                    previous_selector = None
                    while parent and wile_loop_limit < 5:
                        parent_selectors = self.get_selectors(parent)

                        for parent_selector in parent_selectors:

                            if previous_selector is None:
                                current_selector = f"{parent_selector} {selector}"
                            elif previous_selector:
                                current_selector = f"{parent_selector} {previous_selector}"

                            if current_selector in rule.selectorText:
                                if current_selector == rule.selectorText:
                                    self.tag_rules.add(rule)
                                previous_selector = current_selector

                            parent = parent.parent
                            wile_loop_limit += 1

    def get_selectors(self, tag):
        selectors = tag.get('class', [])
        if selectors:
            selectors = [f'.{selector}' for selector in selectors]
        if tag.name:
            selectors = [tag.name] + selectors
        return selectors


class PageStyleProcesser(BaseStyleProcesser):
    def __init__(self, css_styles, document):
        super().__init__(css_styles)
        self.document = document

    def apply_page_margin(self):
        for rule in self.page_styles:
            if "margin" in rule.style.keys():
                # TODO add marging processing if > 1 arguments provided
                if len(rule.style["margin"].split()) == 1:
                    margin = rule.style["margin"]
                    margin = ''.join(re.findall('[^cm]', margin))
                    margin = float(margin)

                    sections = self.document.sections
                    for section in sections:
                        section.top_margin = Cm(margin)
                        section.bottom_margin = Cm(margin)
                        section.left_margin = Cm(margin)
                        section.right_margin = Cm(margin)

    def apply_page_font(self):
        for rule in self.page_styles:
            if "font-size" in rule.style.keys():
                font_size = rule.style["font-size"]
                font_size = ''.join(re.findall('[^px]', font_size))
                font_size_px = int(font_size)
                font_size_pt = convert_px_to_pt(font_size_px)
                style = self.document.styles['Normal']
                font = style.font
                font.size = Pt(font_size_pt)
            if "font-family" in rule.style.keys():
                font_type = rule.style['font-family'].split()[0].strip(',')
                style = self.document.styles['Normal']
                font = style.font
                font.name = font_type.capitalize()

    def apply(self):
        self.apply_page_font()
        self.apply_page_margin()


class ParagraphStyleProcesser(BaseStyleProcesser):
    def __init__(self, tag, css_styles, paragraph):
        super().__init__(css_styles, tag)
        self.paragraph = paragraph

    def recursive_paragraph_processer(self, element):
        if not isinstance(element, bs4.element.Tag):
            run = self.paragraph.add_run(element)
            self.set_font_style(run)
        else:
            paragraph_processer = ParagraphStyleProcesser(element, self.css_styles, self.paragraph)
            element = '\n' + element.contents[0]
            paragraph_processer.recursive_paragraph_processer(element)

    def process(self):
        self.set_aligin()
        for element in self.tag.contents:
            if element != '\n':
                self.recursive_paragraph_processer(element)

    def set_aligin(self):
        for rule in self.tag_rules:
            if 'text-align' in rule.style.keys():
                alignment_type = rule.style['text-align'].upper()
                alignment = getattr(WD_ALIGN_PARAGRAPH, f'{alignment_type}', 'JUSTIFY')
                self.paragraph.alignment = alignment

    def set_font_style(self, run):
        for rule in self.tag_rules:
            if "font-size" in rule.style.keys():
                font_size = rule.style["font-size"]
                font_size = ''.join(re.findall('[^px]', font_size))
                font_size_px = int(font_size)
                font_size_pt = convert_px_to_pt(font_size_px)
                font = run.font
                font.size = Pt(font_size_pt)
            if "font-family" in rule.style.keys():
                font_type = rule.style['font-family'].split()[0].strip(',')
                font = run.font
                font.name = font_type.capitalize()
            if "font-weight" in rule.style.keys() or self.tag.name == 'b':
                font_wight = rule.style['font-weight']
                if font_wight == "bold" or self.tag.name == 'b':
                    run.bold = True


class CellStyleProcesser(BaseStyleProcesser):
    def __init__(self, tag, css_styles, cell):
        super().__init__(css_styles, tag)
        self.cell = cell

    def set_cell_border(self, cell, **kwargs):
        """
        Set cell`s border
        Usage:

        self.set_cell_border(
            cell,
            top={"sz": 12, "val": "single"},
            bottom={"sz": 12, "val": "single"},
            start={"sz": 24, "val": "thick"},
            end={"sz": 12, "val": "dashed"}
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        for edge in ('start', 'top', 'end', 'bottom'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                for key in ["sz", "val"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def apply_column_width(self):
        for docx_cell in self.cell.docx_cells:
            for rule in self.tag_rules:
                if "width" in rule.style.keys():
                    cell_width_text = rule.style["width"]
                    cell_width_px = ''.join(re.findall('[^px]', cell_width_text))
                    cell_width_px = int(cell_width_px)
                    cell_width_pt = convert_px_to_pt(cell_width_px)
                    docx_cell.width = Pt(cell_width_pt)

    def apply_borders(self):
        for docx_cell in self.cell.docx_cells:
            for rule in self.tag_rules:
                if "border-bottom" in rule.style.keys():
                    if 'solid' in rule.style["border-bottom"]:
                        bottom = {"val": "thick", "sz": 2}
                        self.set_cell_border(
                            docx_cell,
                            bottom=bottom
                            )

                if "border" in rule.style.keys():
                    if rule.style.border == 'none':
                        try:
                            bottom
                        except NameError:
                            bottom = None
                        if bottom:
                            self.set_cell_border(
                                docx_cell,
                                start={"val": "nil"},
                                top={"val": "nil"},
                                end={"val": "nil"}
                                )
                        if bottom is None:
                            self.set_cell_border(
                                docx_cell,
                                start={"val": "nil"},
                                top={"val": "nil"},
                                end={"val": "nil"},
                                bottom={"val": "nil"}
                                )

    def apply(self):
        self.apply_borders()
        self.apply_column_width()


class HTML2DOCX:
    def __init__(self, template):
        self.document = self.create()
        self.html_doc = self.get_html_doc(template)
        self.paragraph_tags = {'div', 'p', 'b',  'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}
        self.raw_styles = self.html_doc.find('style').contents[0].strip()

    @property
    def root_tags(self):
        return self.html_doc.find_all(recursive=False)

    def create(self):
        return Document()

    def get_html_doc(self, template):
        html_doc = BeautifulSoup(template, 'html.parser')
        return html_doc

    def set_page_styles(self):
        styler = PageStyleProcesser(self.raw_styles, self.document)
        styler.apply()

    def get_styler(self):
        return PageStyleProcesser(self.raw_styles, self.document)

    def write_docx(self):
        self.set_page_styles()

        for tag in self.root_tags:
            if tag.name in self.paragraph_tags:
                handle_tag = self.process_paragraph_tag
                handle_tag(tag)
            else:
                handle_tag = getattr(self, f'process_{tag.name}_tag', None)
                if handle_tag:
                    handle_tag(tag)

        with tempfile.TemporaryDirectory() as tmpdirname:
            docx_file_path = f'{tmpdirname}/temp_docx.docx'
            self.document.save(docx_file_path)
            with open(docx_file_path, 'rb') as f:
                docx_bytes = f.read()
        return docx_bytes

    def process_paragraph_tag(self, tag):
        paragraph = self.document.add_paragraph()
        paragraph_styler = ParagraphStyleProcesser(tag,
                                                   self.raw_styles,
                                                   paragraph)
        paragraph_styler.process()

    def process_table_tag(self, tag):
        drawer = TableDrawer(tag, self.document, self.raw_styles)
        drawer.draw()
